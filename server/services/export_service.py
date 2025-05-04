import os
import re
import markdown
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import html


class ExportService:
    def __init__(self):
        pass

    def markdown_to_pdf(self, markdown_content, output_file):
        """Convert markdown content to PDF with better formatting"""
        # Process the markdown content to handle formatting
        processed_content = self._preprocess_markdown(markdown_content)

        # Convert markdown to HTML
        html_content = markdown.markdown(processed_content, extensions=[
                                         'tables', 'fenced_code'])

        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Add a title
        pdf.set_font("Arial", 'B', size=16)
        pdf.cell(0, 10, "Documentation", ln=True, align='C')
        pdf.ln(10)

        # Set normal font for content
        pdf.set_font("Arial", size=11)

        # Process HTML content
        html_elements = re.split(
            r'<h[1-4]>|</h[1-4]>|<p>|</p>|<pre>|</pre>|<code>|</code>|<strong>|</strong>|<em>|</em>', html_content)
        current_tag = None

        for element in html_elements:
            if not element.strip():
                continue

            # Check if it's a header
            if re.match(r'<h[1-4]>', element):
                current_tag = element
                pdf.set_font("Arial", 'B', size=14)
                continue
            elif re.match(r'</h[1-4]>', element):
                current_tag = None
                pdf.set_font("Arial", size=11)
                continue

            # Check if it's a paragraph or other tag
            if current_tag and re.match(r'<h[1-4]>', current_tag):
                # It's a header text
                pdf.set_font("Arial", 'B', size=14)
                pdf.multi_cell(0, 10, txt=html.unescape(element.strip()))
                pdf.ln(5)
                pdf.set_font("Arial", size=11)
            else:
                # It's regular content
                pdf.multi_cell(0, 7, txt=html.unescape(element.strip()))
                pdf.ln(2)

        # Add footer with page numbers
        pdf.set_auto_page_break(False)
        for page_num in range(1, pdf.page_no() + 1):
            pdf.set_y(-15)
            pdf.set_font('Arial', 'I', 8)
            pdf.cell(0, 10, f'Page {page_num}', 0, 0, 'C')

        pdf.output(output_file)
        return output_file

    def markdown_to_docx(self, markdown_content, output_file):
        """Convert markdown content to Word document with better formatting"""
        # Process the markdown content to handle formatting
        processed_content = self._preprocess_markdown(markdown_content)

        doc = Document()

        # Add document title
        title = doc.add_heading('Documentation', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a subtle line below the title
        border = doc.add_paragraph()
        border.add_run('_' * 50)
        border.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Split the content into paragraphs
        paragraphs = processed_content.split('\n\n')

        for paragraph in paragraphs:
            # Handle headers
            if paragraph.startswith('# '):
                doc.add_heading(paragraph[2:], level=1)
            elif paragraph.startswith('## '):
                doc.add_heading(paragraph[3:], level=2)
            elif paragraph.startswith('### '):
                doc.add_heading(paragraph[4:], level=3)
            elif paragraph.startswith('#### '):
                doc.add_heading(paragraph[5:], level=4)
            # Handle code blocks
            elif paragraph.startswith('```') and paragraph.endswith('```'):
                code_block = paragraph[3:-3].strip()
                p = doc.add_paragraph()
                code_run = p.add_run(code_block)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                p.style = 'Normal'
                doc.add_paragraph()  # Add some space after code blocks
            # Handle bold and italic text in regular paragraphs
            else:
                p = doc.add_paragraph()

                # Handle inline formatting
                current_pos = 0
                # Match patterns like **bold** or *italic*
                for match in re.finditer(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', paragraph):
                    # Add any text before the match
                    if match.start() > current_pos:
                        p.add_run(paragraph[current_pos:match.start()])

                    # Process the match
                    text = match.group()
                    if text.startswith('**') and text.endswith('**'):
                        # Bold text
                        run = p.add_run(text[2:-2])
                        run.bold = True
                    elif text.startswith('*') and text.endswith('*'):
                        # Italic text
                        run = p.add_run(text[1:-1])
                        run.italic = True
                    elif text.startswith('`') and text.endswith('`'):
                        # Code text
                        run = p.add_run(text[1:-1])
                        run.font.name = 'Courier New'

                    current_pos = match.end()

                # Add any remaining text
                if current_pos < len(paragraph):
                    p.add_run(paragraph[current_pos:])

        # Add footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_text = footer.paragraphs[0]
        footer_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_text.text = "Page "
        run = footer_text.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(9)

        doc.save(output_file)
        return output_file

    def _preprocess_markdown(self, content):
        """Process markdown content to ensure proper format rendering"""
        # Replace **text** that appears at the beginning of a line to ensure proper rendering
        processed_content = re.sub(
            r'^\*\*([^*]+)\*\*', r'**\1**', content, flags=re.MULTILINE)

        # Ensure headers have a space after the # symbols
        processed_content = re.sub(
            r'^(#+)([^#\s])', r'\1 \2', processed_content, flags=re.MULTILINE)

        # Ensure code blocks have proper spacing
        processed_content = re.sub(
            r'```(\w+)?\n', r'```\1\n', processed_content)

        return processed_content

    def export_documentation(self, content, format_type, output_path):
        """Export documentation to the specified format"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        if format_type.lower() == 'pdf':
            return self.markdown_to_pdf(content, output_path)
        elif format_type.lower() == 'docx':
            return self.markdown_to_docx(content, output_path)
        else:
            raise ValueError(
                f"Unsupported format: {format_type}. Supported formats: pdf, docx")
